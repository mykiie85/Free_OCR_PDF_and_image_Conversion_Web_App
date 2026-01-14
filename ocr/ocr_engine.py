import os
import pytesseract
from PIL import Image
import pdf2image
import cv2
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

from .preprocess import ImagePreprocessor
from .layout_parser import LayoutParser


class OCREngine:
    def __init__(self, tesseract_path=None):
        
        self.preprocessor = ImagePreprocessor()
        self.layout_parser = LayoutParser()
        
        # Set Tesseract path
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        else:
            self._auto_detect_tesseract()
    
    def _auto_detect_tesseract(self):
        
        # Common paths
        paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            '/usr/bin/tesseract',
            '/usr/local/bin/tesseract',
        ]
        
        # Try environment variable
        env_path = os.environ.get('TESSERACT_PATH')
        if env_path:
            paths.insert(0, env_path)
        
        for path in paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"✓ Tesseract found: {path}")
                return
        
        print("⚠ Tesseract not found. Please install or set TESSERACT_PATH environment variable")
    
    def check_tesseract(self):
        
        try:
            version = pytesseract.get_tesseract_version()
            return {'available': True, 'version': str(version)}
        except:
            return {'available': False, 'version': None}
    
    def process_document(self, input_path, output_format='txt', output_folder='outputs', file_id='', language='eng'):
        
        try:
            print(f"Processing: {input_path} with language: {language}")
            
            # Check file exists
            if not os.path.exists(input_path):
                return {'success': False, 'error': 'File not found'}
            
            # Convert to images
            images = self._convert_to_images(input_path)
            
            if not images:
                return {'success': False, 'error': 'Failed to convert file to images'}
            
            # Check page limit
            if len(images) > 50:
                return {'success': False, 'error': f'Too many pages ({len(images)}). Maximum is 50 pages.'}
            
            print(f"Processing {len(images)} page(s)...")
            
            # Extract structured data from all pages
            pages_data = []
            for i, image in enumerate(images, 1):
                print(f"OCR on page {i}/{len(images)}...")
                page_data = self._extract_page_data(image, i, language)
                pages_data.append(page_data)
            
            # Generate output based on format
            output_path = self._generate_output(
                pages_data=pages_data,
                output_format=output_format,
                output_folder=output_folder,
                original_filename=Path(input_path).stem,
                file_id=file_id
            )
            
            return {
                'success': True,
                'output_path': output_path,
                'pages': len(images)
            }
        
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _convert_to_images(self, input_path):
        
        try:
            file_ext = Path(input_path).suffix.lower()
            
            if file_ext == '.pdf':
                # Convert PDF to images
                images = pdf2image.convert_from_path(input_path, dpi=300)
                return images
            else:
                # Load single image
                image = Image.open(input_path)
                return [image]
        
        except Exception as e:
            print(f"Error converting to images: {e}")
            return []
    
    def _extract_page_data(self, pil_image, page_num, language='eng'):
        
        # Convert PIL to OpenCV
        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
        
        # Resize for optimal OCR
        cv_image = self.preprocessor.resize_for_ocr(cv_image)
        
        # Preprocess with advanced method
        processed = self.preprocessor.preprocess(cv_image, method='advanced')
        
        # Convert back to PIL
        pil_processed = Image.fromarray(processed)
        
        # Use better Tesseract config for higher accuracy
        # PSM 3 = Fully automatic page segmentation (better for documents)
        # OEM 3 = Default OCR Engine Mode (LSTM neural networks)
        custom_config = r'--oem 3 --psm 3'
        
        # Extract text with enhanced config
        text = pytesseract.image_to_string(pil_processed, lang=language, config=custom_config)
        
        # Get detailed data for layout with same config
        data = pytesseract.image_to_data(
            pil_processed, 
            lang=language, 
            config=custom_config,
            output_type=pytesseract.Output.DICT
        )
        
        # Detect tables
        tables = self.layout_parser.detect_tables(cv_image)
        
        # Parse layout structure
        blocks = self.layout_parser.parse_layout(data)
        
        return {
            'page_num': page_num,
            'text': text,
            'blocks': blocks,
            'tables': tables,
            'data': data
        }
    
    def _generate_output(self, pages_data, output_format, output_folder, original_filename, file_id):
        
        
        # Create output filename
        output_filename = f"{file_id}_{original_filename}.{output_format}"
        output_path = os.path.join(output_folder, output_filename)
        
        if output_format == 'txt':
            self._generate_txt(pages_data, output_path)
        elif output_format == 'docx':
            self._generate_docx(pages_data, output_path)
        elif output_format == 'xlsx':
            self._generate_xlsx(pages_data, output_path)
        
        return output_path
    
    def _generate_txt(self, pages_data, output_path):
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for page_idx, page in enumerate(pages_data):
                # Add page separator (except for first page)
                if page_idx > 0:
                    f.write(f"\n{'='*60}\n")
                    f.write(f"PAGE {page['page_num']}\n")
                    f.write(f"{'='*60}\n\n")
                
                # Clean up text
                text = page['text'].strip()
                
                # Remove excessive blank lines while preserving paragraph structure
                lines = text.split('\n')
                cleaned_lines = []
                prev_blank = False
                
                for line in lines:
                    stripped = line.strip()
                    
                    if stripped:
                        # Add the line with proper indentation preserved
                        cleaned_lines.append(line.rstrip())
                        prev_blank = False
                    elif not prev_blank:
                        # Add single blank line for paragraph separation
                        cleaned_lines.append('')
                        prev_blank = True
                
                # Remove trailing blank lines
                while cleaned_lines and not cleaned_lines[-1]:
                    cleaned_lines.pop()
                
                # Write cleaned text
                f.write('\n'.join(cleaned_lines))
                f.write('\n\n')
        
        print(f"✓ TXT saved: {output_path}")
    
    def _generate_docx(self, pages_data, output_path):
        
        doc = Document()
        
        # Set default font and spacing
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        for page_idx, page in enumerate(pages_data):
            # Don't add page heading for first page or if text looks like a title
            if page_idx > 0:
                doc.add_heading(f'Page {page["page_num"]}', level=2)
            
            # Get all text lines for better paragraph detection
            full_text = page['text']
            lines = [line.strip() for line in full_text.split('\n') if line.strip()]
            
            current_paragraph = []
            
            for line in lines:
                # Detect headings (all caps, short, or ends with specific patterns)
                is_heading = (
                    (len(line) < 80 and line.isupper()) or
                    (len(line) < 50 and line.endswith(':')) or
                    line.startswith('YAH:') or
                    line.startswith('REF:') or
                    line.startswith('TAREHE:') or
                    line.startswith('Kumb.')
                )
                
                # Check if it's a new paragraph (blank line indicator in structure)
                # Or if it starts with a number followed by period (numbered item)
                is_new_paragraph = (
                    line and len(line) > 2 and 
                    (line[0].isdigit() and line[1] in ['.', ')'])
                )
                
                if is_heading:
                    # Flush current paragraph if exists
                    if current_paragraph:
                        para_text = ' '.join(current_paragraph)
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.15
                        current_paragraph = []
                    
                    # Add as heading
                    if line.isupper() and len(line) < 80:
                        heading = doc.add_heading(line, level=3)
                        heading.paragraph_format.space_before = Pt(12)
                        heading.paragraph_format.space_after = Pt(6)
                    else:
                        p = doc.add_paragraph(line)
                        run = p.runs[0]
                        run.bold = True
                        p.paragraph_format.space_after = Pt(8)
                
                elif is_new_paragraph:
                    # Flush current paragraph
                    if current_paragraph:
                        para_text = ' '.join(current_paragraph)
                        p = doc.add_paragraph(para_text)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.15
                        current_paragraph = []
                    
                    # Start new numbered paragraph
                    current_paragraph = [line]
                
                else:
                    # Continue building paragraph
                    current_paragraph.append(line)
            
            # Flush any remaining paragraph
            if current_paragraph:
                para_text = ' '.join(current_paragraph)
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
            
            # Add tables if detected
            for table_data in page['tables']:
                try:
                    if table_data.get('rows') and len(table_data['rows']) > 0:
                        # Get max number of columns
                        max_cols = max(len(row) for row in table_data['rows']) if table_data['rows'] else 0
                        
                        if max_cols > 0:
                            # Add spacing before table
                            doc.add_paragraph()
                            
                            table = doc.add_table(rows=len(table_data['rows']), cols=max_cols)
                            table.style = 'Light Grid Accent 1'
                            
                            for i, row_data in enumerate(table_data['rows']):
                                for j, cell_text in enumerate(row_data):
                                    if j < max_cols:
                                        table.rows[i].cells[j].text = str(cell_text)
                            
                            # Add spacing after table
                            doc.add_paragraph()
                except Exception as e:
                    print(f"Error adding table: {e}")
                    continue
            
            # Page break (except last page)
            if page != pages_data[-1]:
                doc.add_page_break()
        
        doc.save(output_path)
        print(f"✓ DOCX saved: {output_path}")
    
    def _generate_xlsx(self, pages_data, output_path):
        
        try:
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                for page in pages_data:
                    sheet_name = f"Page_{page['page_num']}"
                    
                    # If tables detected, write them
                    if page['tables']:
                        row_offset = 0
                        for table_idx, table in enumerate(page['tables']):
                            try:
                                if table.get('rows') and len(table['rows']) > 0:
                                    # Get max columns to handle uneven rows
                                    max_cols = max(len(row) for row in table['rows'])
                                    
                                    # Pad rows to have equal length
                                    padded_rows = []
                                    for row in table['rows']:
                                        padded_row = list(row) + [''] * (max_cols - len(row))
                                        padded_rows.append(padded_row)
                                    
                                    df = pd.DataFrame(padded_rows)
                                    df.to_excel(
                                        writer, 
                                        sheet_name=sheet_name, 
                                        startrow=row_offset, 
                                        index=False, 
                                        header=False
                                    )
                                    row_offset += len(padded_rows) + 2
                            except Exception as e:
                                print(f"Error writing table to Excel: {e}")
                                continue
                    else:
                        # No tables - write text as single column
                        try:
                            text_lines = [line.strip() for line in page['text'].split('\n') if line.strip()]
                            if text_lines:
                                df = pd.DataFrame({'Content': text_lines})
                                df.to_excel(writer, sheet_name=sheet_name, index=False)
                        except Exception as e:
                            print(f"Error writing text to Excel: {e}")
                            # Create empty sheet
                            df = pd.DataFrame({'Content': ['No content extracted']})
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            print(f"✓ Excel saved: {output_path}")
        except Exception as e:
            print(f"Error generating Excel file: {e}")
            # Create a simple fallback file
            df = pd.DataFrame({'Error': ['Failed to generate Excel file']})
            df.to_excel(output_path, index=False)